#!/usr/bin/env python3
# ============================================================
# 各プロダクトの「事業計画書」＋「技術仕様書」を Word(.docx) で生成
# 出力: ~/Desktop/CORE 営業資料/<Product>/CORE_<Product>_事業計画書.docx 等
# ============================================================
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.expanduser('~/Desktop/CORE 営業資料')
NAVY = RGBColor(0x0F, 0x1B, 0x33)
ACC  = RGBColor(0x2B, 0x54, 0xE6)
GRAY = RGBColor(0x5A, 0x64, 0x78)

def new_doc(title, subtitle, accent):
    d = Document()
    st = d.styles['Normal']; st.font.name = 'Hiragino Sans'; st.font.size = Pt(10.5)
    # title block
    p = d.add_paragraph(); r = p.add_run('株式会社コア — CORE Inc.'); r.font.size=Pt(9); r.font.color.rgb=GRAY
    h = d.add_paragraph(); r = h.add_run(title); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
    s = d.add_paragraph(); r = s.add_run(subtitle); r.font.size=Pt(11); r.font.color.rgb=RGBColor(*accent)
    d.add_paragraph('作成日: 2026-06-10 ／ 機密（社外秘）').runs[0].font.size=Pt(8)
    d.add_paragraph()
    return d

def H(d, text, accent):
    p = d.add_paragraph(); p.space_before=Pt(10)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor(*accent)
    return p

def H2(d, text):
    p = d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=NAVY; return p

def P(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.font.size=Pt(10.5); return p

def BUL(d, items):
    for it in items:
        p=d.add_paragraph(style='List Bullet'); r=p.add_run(it); r.font.size=Pt(10.5)

def TABLE(d, headers, rows):
    t=d.add_table(rows=1, cols=len(headers)); t.style='Light Grid Accent 1'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; r=cells[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(9.5)
    d.add_paragraph()

# ============================================================
PRODUCTS = {
'Prism': dict(acc=(0x4B,0x57,0xB0), role='全事業の司令塔・経営AIエージェントOS', url='core-prism-app.vercel.app',
  biz=dict(
   summary='Prism は、経営者の中に同居する複数の役割（決断・営業・財務・創造）を、専属のAIエージェントとして外部化する「経営OS」。13名のAI役員が並走し、議事録・提案・契約・財務などの実務を下書きまで自動化することで、経営者の時間を「作業」から「判断」へと取り戻す。',
   problem=['経営者の時間が、議事録・提案書・契約レビュー・経理・日程調整などの作業に奪われ、本来の意思決定に向き合えない。',
            '顧問・税理士・秘書・各種SaaSを個別契約すると月数十万円、かつ窓口が分散し管理コストが高い。',
            '商談・数字・契約・ナレッジが各ツールに散在し、必要なときに探せず意思決定が遅れる。'],
   market='国内の個人事業主・中小企業経営者は約数百万。生成AIの業務活用は黎明期で、「単体ツール」は普及しつつも「現場の成果まで担うAI」は空白。判断業務の代替価値（コンサル/秘書/専門士業）で課金できる点が差別化。',
   target=['一人〜数十名の経営者・個人事業主','士業（税理士・弁護士・行政書士）・コンサルタント','副業/複業のマルチロール経営者・スタートアップ'],
   pricing=[['7日間無料','¥0/7日間','全機能トライアル・クレカ不要'],['Starter','¥4,800/月','個人・基本AI'],
            ['Standard','¥9,800/月','全AI・無制限（人気No.1）'],['Exclusive','¥29,800/月','法人・専任CS・導入伴走']],
   diff='「AIツール」ではなく「AI役員」。操作させるのではなく、任せて確認する。判断業務（コンサル/秘書/士業）の代替価値で課金し、複数SaaSをひとつのOSに置換する点が堀。',
   revenue='サブスクリプション（月/年）。年額は月額×10（約17%off）。上位プランほどLTVが高く、法人プランで請求書払い・チーム導入を取り込む。AI原価はモデル選択（Haiku中心）で最適化し高粗利を狙う。',
   roadmap=[['2026','4プロダクト本格ローンチ・個人/中小向け獲得'],['2027','法人プラン強化（チーム・請求書払い・SSO）'],['2028〜','英/韓/台展開、東アジアの中小経営者へ']],
   risks=['生成AIの価格/性能変動 → モデル抽象化レイヤとBYOKで吸収','大手SaaSのAI内製化 → 「役割分担×一気通貫×日本の中小特化」で差別化','品質/ハルシネーション → 承認制・根拠提示・honest-numbers原則']),
  tech=dict(
   overview='Vite + React + TypeScript の SPA（マルチページ：/ Prism, /iris, /corp）。Vercel にホスティングし、サーバーレス関数（api/）で AI・Stripe・各種連携を処理。状態はブラウザ localStorage を基本に、課金・同期はサーバー真実で担保。',
   arch=['フロント: Vite + React 19 + TypeScript（コード分割・遅延ロード）','配信: Vercel（CDN + サーバーレス関数 api/）',
         'AI: Anthropic Claude（Haiku中心）/ Google Gemini を用途別に使い分け','課金: Stripe（Payment Links + Checkout、billing.ts でプラン定義）',
         '永続化: ブラウザ localStorage（作業状態）＋ サーバーレスでの整合確認'],
   stack=[['フロント','Vite, React 19, TypeScript, framer-motion'],['UI','独自デザインシステム（Tailwind一部）'],
          ['AI','Anthropic SDK（Claude）, Gemini'],['課金','Stripe（live）, Payment Links, Webhook'],['配信','Vercel（api/ serverless）']],
   features=[['7役割×AIエージェント','ペルソナ単位で文脈・ナレッジを分離'],['議事録/提案/契約レビュー','音声/ファイル→Claudeで下書き生成'],
             ['横断検索','全文脈の統合検索'],['能動提案','朝晩の巡回でタスク提案']],
   data='ユーザーの作業状態・ナレッジは localStorage（端末内）。課金状態は Stripe をソース・オブ・トゥルースとし、entitlement をサーバーで判定。',
   integrations=['Stripe（決済・サブスク）','Google（カレンダー/Gmail 連携）','Anthropic / Gemini（AI）'],
   security=['課金はStripe live鍵をサーバー側のみ保持（クライアント非公開）','CSP/HSTS等のセキュリティヘッダ適用','個人データは端末内中心・最小権限'],
   billing='Stripe の Payment Links（VITE_STRIPE_*_URL）と Checkout を併用。全プラン subscription_data.trial_period_days=7（8日目課金）。プラン定義は billing.ts。',
   infra='Vercel（プレビュー/本番）。git連携の自動デプロイ＋CLI強制デプロイ。',
   nfr=['初回表示の軽量化（チャンク分割）','7日間トライアルはカード登録不要','多言語（日/英/中）対応の i18n 基盤'])),

'Iris': dict(acc=(0xA8,0x49,0x7B), role='インフルエンサーのためのInstagram運用AI', url='core-prism-app.vercel.app/iris',
  biz=dict(
   summary='Iris は、Instagram運用の6業務（投稿制作・分析・案件管理・DM返信・交渉・美容相談）をひとつのAIアプリに束ねる。数字を見せるだけの分析ツールと異なり、解析を踏まえて「次の一手」まで提案し、「いいね」を「案件」に変える。',
   problem=['投稿・分析・案件・DM・交渉を一人で抱え、創作時間が削られる。','フォロワーは増えても案件/売上につながらない（数字と稼ぎの乖離）。','既存分析ツールは数字を見せるだけで、打ち手を示さない。'],
   market='国内のInstagramクリエイター/インフルエンサーは多数。分析特化（トレミル等）や法人向け高額ツール（SINIS等）は存在するが、「案件管理＋交渉＋投稿制作＋次の一手提案」を個人価格で束ねる製品は空白。',
   target=['個人インフルエンサー・クリエイター','SNS運用代行・マネージャー','事務所/代理店（上位プラン）'],
   pricing=[['7日間無料','¥0/7日間','全機能トライアル・クレカ不要'],['Lite','¥2,980/月','入門・副業'],['Standard','¥6,980/月','本気のクリエイター（人気No.1）'],
            ['Pro','¥12,800/月','チーム/代理店・運用代行']],
   diff='数字を見せるだけでなく「次の一手」まで決める。案件管理（スクショ3秒入力）・交渉文AI・投稿制作を統合。司令塔Prismや配信Resonanceと連携。',
   revenue='サブスク（月/年）。Standard を主力に、Pro/Studioでチーム・代理店を取り込む。BYOK的なAI原価最適化で高粗利。',
   roadmap=[['2026','個人クリエイター獲得・コミュニティ形成'],['2027','ブランドマッチ（企業×クリエイター双方向）強化'],['2028〜','多言語・海外クリエイター']],
   risks=['Instagram API/規約変更 → 公式API準拠と手動補完','競合の機能追随 → 一気通貫とエコシステム連携で差別化']),
  tech=dict(
   overview='Prism と同一コードベース（core-prism-app）の /iris ルートとして提供。Vite+React+TS、Vercel配信。Instagram解析・投稿生成・案件管理をAIで実装。',
   arch=['フロント: Vite + React + TypeScript（/iris ルート）','AI: Claude（文章/戦略）・画像/動画系AI（投稿素材）','課金: Stripe Payment Links（VITE_STRIPE_IRIS_*_URL）','状態: localStorage＋entitlement判定'],
   stack=[['フロント','Vite, React, TypeScript'],['AI','Anthropic Claude, Gemini, 画像生成'],['課金','Stripe（live, IRIS_*）'],['配信','Vercel']],
   features=[['投稿AI','構成/テロップ/キャプション/タグ生成'],['Instagram解析','指標取得→次の一手提案'],['案件管理','スクショ→AI自動入力'],['交渉文AI','返信/断り/カウンター生成']],
   data='案件・履歴は端末/サーバーに保持。課金状態は Stripe を真実とする。',
   integrations=['Instagram（解析）','Stripe','Anthropic/Gemini','Instagramシェア（Web Share/URL Scheme）'],
   security=['APIキーはサーバー側保持','個人/案件データの最小権限管理'],
   billing='Stripe Payment Links（Lite/Standard/Pro/Studio）＋全プラン7日間トライアル（trial_period_days=7）。',
   infra='Vercel（core-prism-app に同梱）。',
   nfr=['モバイルPWA最適化','投稿素材生成の応答速度'])),

'Resonance': dict(acc=(0x2E,0x8B,0x6F), role='店舗・サロン・教室のためのLINE個別配信SaaS', url='resonancebot-ivory.vercel.app',
  biz=dict(
   summary='Resonance は、名簿の一人ひとりにAIが文面を書き分け、LINEで手紙のように届ける個別配信SaaS。承認制で全件確認でき、既存のLINE公式アカウントに接続。BYOK（各社の鍵）でAI原価をほぼ0にし、低価格を実現。',
   problem=['一斉配信は全員同じ文面で刺さらず、ブロック/離脱を招く。','店舗の生命線である再来店（LTV）が、一度きりの来店で途切れる。','一人ひとりへの個別対応は手間が大きく、現場では続かない。'],
   market='国内の小規模事業者は約285万、LINE公式アカウント活用は数十万規模。運用課題（パーソナライズ・継続）は大きい。競合（LINE公式/Lステップ/エルメ）は「セグメントに同じ文」止まりで、AIの個別文面＋承認制は空白。',
   target=['店舗・サロン・スクール（教室）','アーティスト/インフルエンサーのファン運用','LINE公式アカウントを持つ小規模事業者'],
   pricing=[['7日間無料','¥0/7日間','全機能トライアル・クレカ不要'],['Solo','¥1,980/月','1アカウント・月2,000通'],
            ['Pro','¥4,980/月','AIレター・月8,000通（人気No.1）'],['Business','¥9,800/月','3アカウント・月30,000通・設定代行']],
   diff='セグメント配信（同じ文）ではなく、一人ひとり別文面＋送信前の全件確認。BYOKでAI原価ほぼ0＝価格の堀。Iris/Lumeの来訪データで宛先最適化。',
   revenue='サブスク（月）。BYOKによりAI原価が極小で高粗利（粗利率96%級）。アカウント数/通数で上位プランへ拡張。参考：TAM約32.6万件／SAM約13万／SOM約2,600件・LTV/CAC約11倍（自社試算）。',
   roadmap=[['2026','店舗/サロン/教室への展開・実LINE連携拡大'],['2027','セグメント提案AI・予約/CRM連携'],['2028〜','多店舗/チェーン・パートナー販売']],
   risks=['LINEの規約/料金変更 → 公式API準拠・BYOK設計','個人LINEでの一斉送信は規約違反/凍結リスク → 公式アカウント誘導を正攻法として徹底']),
  tech=dict(
   overview='Next.js（App Router）+ TypeScript の SaaS。Vercel配信。Upstash Redis（KV）で名簿/履歴をマルチテナント分離。LINE Messaging API と Stripe を統合。BYOK（各社のLINE/Claude鍵）。',
   arch=['フロント/サーバー: Next.js 16（Route Handlers, ステートレスAPI）','永続化: Upstash Redis（REST）。キー名前空間 rb:followers:{basicId} / rb:msgs:{basicId}:{userId}',
         'LINE: Messaging API（push/broadcast/getProfile）、Webhook署名検証','AI: Anthropic Claude（個別文面・AIレター。BYOK）','課金: Stripe（ResonanceBotアカウント, live）'],
   stack=[['フレームワーク','Next.js 16, TypeScript'],['UI','Tailwind 4, shadcn/ui'],['KV','Upstash Redis（REST）'],
          ['メッセージング','LINE Messaging API'],['AI','Anthropic Claude（haiku）'],['課金','Stripe（live, rk_live）'],['配信','Vercel']],
   features=[['個別文面AI','名簿×履歴からClaudeが個別下書き（{name}差込）'],['承認制配信','送信前に全件確認、AIは自動送信しない'],
             ['Webhook収集','follow/messageでgetProfile→KV名簿upsert・署名検証'],['一斉/個別送信','line-broadcast（全員）/ line-send-each（個別）']],
   data='名簿・会話履歴は Upstash Redis にテナント（basicId）別で分離保存。接続トークンはクライアント localStorage に保持（public URLのため送信は入力トークン使用）。',
   integrations=['LINE Messaging API（push/broadcast/webhook）','Stripe（サブスク・トライアル）','Anthropic Claude（BYOK）','QR/友だち招待（公式アカウント誘導）'],
   security=['Webhook署名検証（LINE_CHANNEL_SECRET）。未署名POSTは401','送信APIはクライアント入力トークン使用（envトークンはWebhook自動返信専用＝乱用防止）','テナント分離（basicId名前空間）'],
   billing='Stripe Checkout（ResonanceBotアカウント, rk_live）。STRIPE_PRICE_PRO/BUSINESS/PREMIUM。全有料プラン trial_period_days=7（8日目課金）。confirmでKV rb:plan:{ns} に保存、entitlementでサーバー判定。',
   infra='Vercel（production env に LINE_*/STRIPE_*/Upstash KV を注入）。',
   nfr=['ステートレスAPIでサーバーレス整合','モバイル/PWA最適化','送信前確認による誤送信防止'])),

'Lume': dict(acc=(0x6A,0x57,0xC8), role='クリエイターのためのリンクハブ＋クリック解析', url='lume-deploy-five.vercel.app',
  biz=dict(
   summary='Lume は、全リンクを美しいプロフィールに30秒で集約し、クリックをヒートマップ・流入元・時間帯で可視化するリンクハブ。リンクを「ただの入口」から「測れる資産」へ変え、Iris/Resonance/Prismの判断材料にする。',
   problem=['SNSのプロフィールに置けるリンクは基本ひとつで、複数発信先へ誘導できない。','リンクを並べても誰がどこから何を踏んだか分からず効果が測れない。','各SNS/予約/販売の導線が分断され、成果が追えない。'],
   market='リンクまとめ（Linktree等）は普及するが、「美しさ＋クリック解析（ヒートマップ/流入元/時間帯）」を日本語・低価格で提供する製品は空白。クリエイター/店舗のプロフィール運用が対象。',
   target=['クリエイター・インフルエンサー','店舗/サロン/教室','あらゆる発信者（プロフィールリンク運用者）'],
   pricing=[['7日間無料','¥0/7日間','全機能トライアル・クレカ不要'],['Pro','¥1,480/月','解析フル機能（人気）'],['Business','¥3,480/月','複数プロフィール管理']],
   diff='ただ並べるだけのリンクまとめと違い、クリックを熱で可視化。流入元クロス分析・時間帯傾向まで。エコシステム（Iris/Resonance/Prism）に来訪データを供給。',
   revenue='サブスク（月）。Proを主力に、Businessで複数プロフィール/チーム。低原価・高粗利。',
   roadmap=[['2026','クリエイター/店舗への展開・テーマ拡充'],['2027','EC/予約連携・A/Bテスト'],['2028〜','エコシステム連携の深化']],
   risks=['競合（Linktree等）の機能追随 → 解析の深さとエコシステム連携で差別化','プライバシー/計測規制 → 同意取得・最小データ']),
  tech=dict(
   overview='公開LP/アプリは静的配信（lume-deploy）＋バックエンド（lume-backend）構成。Vercel配信。リンク集約と、クリックの収集・可視化（ヒートマップ/流入元/時間帯）を提供。',
   arch=['フロント: 静的HTML/CSS/JS（高速表示）、テーマ切替','バックエンド: lume-backend（計測収集・集計API）','解析: クリックイベント収集→ヒートマップ/流入元/時間帯に集計','課金: Stripe（Lume Pro/Business, CORE Incアカウント）'],
   stack=[['フロント','静的HTML/CSS/JS（軽量）'],['バックエンド','lume-backend（計測API）'],['解析','クリックイベント集計（ヒートマップ等）'],['課金','Stripe（live, Lume Pro/Business）'],['配信','Vercel']],
   features=[['30秒で集約','5テーマの美しいプロフィール'],['クリックヒートマップ','押下比率を熱で可視化'],['流入元クロス分析','参照元×リンクの分析'],['時間帯・傾向','時間帯別の可視化']],
   data='クリック/来訪イベントを収集・集計。個人特定は行わず、傾向データとして可視化。',
   integrations=['Stripe（サブスク）','各SNS/予約/EC（リンク先）','（連携）Iris/Resonance/Prismへ来訪データ供給'],
   security=['計測は同意ベース・最小データ','課金はStripeサーバー側'],
   billing='Stripe（会社口座, Lume Pro ¥1,480 / Business ¥3,480）。7日間トライアル（trial_period_days=7、8日目課金）に統一。',
   infra='Vercel（lume-deploy 静的 + lume-backend）。',
   nfr=['高速表示（静的配信）','モバイル最適化','計測の正確性'])),
}

def build_biz(name, p):
    d = new_doc(f'{name} 事業計画書', f'CORE {name} — {p["role"]}', p['acc'])
    b = p['biz']
    H(d,'1. エグゼクティブサマリー',p['acc']); P(d,b['summary'])
    H(d,'2. 解決する課題',p['acc']); BUL(d,b['problem'])
    H(d,'3. 市場機会',p['acc']); P(d,b['market'])
    H(d,'4. ターゲット顧客',p['acc']); BUL(d,b['target'])
    H(d,'5. 料金モデル',p['acc']); TABLE(d,['プラン','価格','内容'],b['pricing'])
    P(d,'※ 全プラン7日間無料トライアル・クレジットカード登録不要（8日目以降に課金）。')
    H(d,'6. 差別化・競争優位',p['acc']); P(d,b['diff'])
    H(d,'7. 収益モデル',p['acc']); P(d,b['revenue'])
    H(d,'8. ロードマップ',p['acc']); TABLE(d,['時期','マイルストーン'],b['roadmap'])
    H(d,'9. リスクと対策',p['acc']); BUL(d,b['risks'])
    out=f'{BASE}/{name}/CORE_{name}_事業計画書.docx'; d.save(out); print('wrote',out)

def build_tech(name, p):
    d = new_doc(f'{name} 技術仕様書', f'CORE {name} — Technical Specification', p['acc'])
    t = p['tech']
    H(d,'1. システム概要',p['acc']); P(d,t['overview'])
    H(d,'2. アーキテクチャ',p['acc']); BUL(d,t['arch'])
    H(d,'3. 技術スタック',p['acc']); TABLE(d,['レイヤー','技術'],t['stack'])
    H(d,'4. 主要機能の実装',p['acc']); TABLE(d,['機能','実装概要'],t['features'])
    H(d,'5. データ・ストレージ',p['acc']); P(d,t['data'])
    H(d,'6. 外部連携',p['acc']); BUL(d,t['integrations'])
    H(d,'7. セキュリティ',p['acc']); BUL(d,t['security'])
    H(d,'8. 課金（Stripe）',p['acc']); P(d,t['billing'])
    H(d,'9. インフラ・デプロイ',p['acc']); P(d,t['infra'])
    H(d,'10. 非機能要件',p['acc']); BUL(d,t['nfr'])
    out=f'{BASE}/{name}/CORE_{name}_技術仕様書.docx'; d.save(out); print('wrote',out)

for name,p in PRODUCTS.items():
    build_biz(name,p); build_tech(name,p)
print('ALL DONE')
