#!/usr/bin/env python3
# ============================================================
# CORE Identity OS — サービス機能説明書 (Word)
# 運営元: 株式会社 CORE
# 出力: ~/Desktop/CORE_Identity_OS_機能説明書_2026-05-11.docx
# ============================================================
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── ページ設定 (B5 寄りの製本サイズ、余白多め) ──
section = doc.sections[0]
section.page_width = Cm(21.0)   # A4 縦
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 基本スタイル設定 ──
styles = doc.styles

def set_font(style_name, font_name, size_pt, color=None, bold=False):
    style = styles[style_name]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if color is not None:
        style.font.color.rgb = color
    style.font.bold = bold
    # 日本語フォント (East Asia)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')

set_font('Normal', 'Hiragino Kaku Gothic ProN', 10.5)
set_font('Heading 1', 'Hiragino Mincho ProN', 24, RGBColor(0x0A, 0x0A, 0x14), bold=True)
set_font('Heading 2', 'Hiragino Kaku Gothic ProN', 16, RGBColor(0x00, 0x33, 0xA0), bold=True)
set_font('Heading 3', 'Hiragino Kaku Gothic ProN', 13, RGBColor(0x0A, 0x0A, 0x14), bold=True)
set_font('Heading 4', 'Hiragino Kaku Gothic ProN', 11, RGBColor(0x6E, 0x6E, 0x73), bold=True)

# Normal の行間
n_style = styles['Normal']
n_style.paragraph_format.line_spacing = 1.7
n_style.paragraph_format.space_after = Pt(4)


# ── ヘルパー ──────────────────────────────────────────

def add_para(text, style=None, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text:
        r = p.add_run(text)
        if size:
            r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        r.font.bold = bold
        r.font.italic = italic
        # JP font
        rPr = r._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_h1(text):
    return add_para(text, style='Heading 1')

def add_h2(text):
    return add_para(text, style='Heading 2')

def add_h3(text):
    return add_para(text, style='Heading 3')

def add_h4(text):
    return add_para(text, style='Heading 4')

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.line_spacing = 1.6
    r = p.add_run(text)
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
    rPr.append(rFonts)
    return p

def page_break():
    doc.add_page_break()

def add_horizontal_line():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'FCB045')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_callout_box(title, body, color_hex='F2F2F7'):
    """淡い背景の囲みボックス"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    # セル背景色
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)
    cell.text = ''
    title_p = cell.paragraphs[0]
    r = title_p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0xA0)
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
    rPr.append(rFonts)
    body_p = cell.add_paragraph()
    body_r = body_p.add_run(body)
    body_r.font.size = Pt(10)
    body_r.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
    rPr2 = body_r._r.get_or_add_rPr()
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
    rPr2.append(rFonts2)
    return table

def add_feature_block(name, purpose, how, output_sample, scenes, plan_note):
    """機能の標準ブロック"""
    add_h3(name)
    add_para(purpose, italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8)
    add_para('● 何ができるか', bold=True, color=RGBColor(0x00, 0x33, 0xA0), size=10.5)
    for h in how:
        add_bullet(h)
    add_para('● 出力例 / 入力例', bold=True, color=RGBColor(0x00, 0x33, 0xA0), size=10.5)
    add_callout_box('SAMPLE', output_sample, 'F2F2F7')
    add_para('● 主な活用シーン', bold=True, color=RGBColor(0x00, 0x33, 0xA0), size=10.5)
    for sc in scenes:
        add_bullet(sc)
    add_para('● プラン別の使い方', bold=True, color=RGBColor(0x00, 0x33, 0xA0), size=10.5)
    add_para(plan_note, size=10, color=RGBColor(0x6E, 0x6E, 0x73), space_after=14)
    add_horizontal_line()

def add_plan_table(plans, columns):
    """プラン比較表"""
    t = doc.add_table(rows=1, cols=len(columns))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, c in enumerate(columns):
        hdr[i].text = c
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x00, 0x33, 0xA0)
    for row in plans:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    rPr = r._r.get_or_add_rPr()
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
                    rPr.append(rFonts)


# ═══════════════════════════════════════════════════════════════════
# 表紙
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()
add_para('CORE  Identity  OS', style='Heading 1', align='center', size=36)
add_para('Service Operation Manual', italic=True, color=RGBColor(0x6E, 0x6E, 0x73), align='center', size=14, space_after=6)
add_horizontal_line()
doc.add_paragraph()
add_para('CORE Prism — 事業家のための AI エージェント OS', align='center', size=12, color=RGBColor(0x00, 0x33, 0xA0))
add_para('CORE Iris — 創作者のための AI エージェント OS', align='center', size=12, color=RGBColor(0xE1, 0x30, 0x6C))
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_para('2026 年 5 月 11 日 発行  ·  Version 1.0', align='center', size=10, color=RGBColor(0x6E, 0x6E, 0x73))
doc.add_paragraph()
add_para('運営元 ──── 株式会社 CORE', align='center', size=11, bold=True, color=RGBColor(0x0A, 0x0A, 0x14))
add_para('https://core-prism-app.vercel.app/  ·  hello@core-inc.jp', align='center', size=9, color=RGBColor(0x6E, 0x6E, 0x73))
page_break()


# ═══════════════════════════════════════════════════════════════════
# 目次
# ═══════════════════════════════════════════════════════════════════
add_h1('目次')
doc.add_paragraph()
toc_items = [
    ('はじめに', '4'),
    ('  本マニュアルの使い方', ''),
    ('  CORE Identity OS の全体像', ''),
    ('', ''),
    ('第 1 部 — CORE Prism (事業家向け)', '8'),
    ('  1.1 Prism とは', ''),
    ('  1.2 7 つのエージェント', ''),
    ('    1.2.1 経営エージェント (CEO Agent)', ''),
    ('    1.2.2 営業エージェント (Sales Agent)', ''),
    ('    1.2.3 財務エージェント (CFO Agent)', ''),
    ('    1.2.4 創造エージェント (Creative Agent)', ''),
    ('    1.2.5 学びエージェント (Knowledge Agent)', ''),
    ('    1.2.6 人材エージェント (People Agent)', ''),
    ('    1.2.7 生活エージェント (Life Agent)', ''),
    ('  1.3 Prism 横断機能 (20 機能)', ''),
    ('  1.4 Prism のプラン詳細', ''),
    ('', ''),
    ('第 2 部 — CORE Iris (創作者向け)', '36'),
    ('  2.1 Iris とは', ''),
    ('  2.2 6 つのファセット', ''),
    ('    2.2.1 案件 (Briefs)', ''),
    ('    2.2.2 分析 (Analytics)', ''),
    ('    2.2.3 創作 (Creation)', ''),
    ('    2.2.4 交渉 (Negotiation)', ''),
    ('    2.2.5 ブランド (Brand)', ''),
    ('    2.2.6 仲間 (Community)', ''),
    ('  2.3 Iris 横断機能 (20 機能)', ''),
    ('  2.4 Iris のプラン詳細', ''),
    ('', ''),
    ('第 3 部 — 両ブランド共通機能', '62'),
    ('  3.1 PWA / オフライン', ''),
    ('  3.2 マルチテナント / 認証', ''),
    ('  3.3 多言語対応', ''),
    ('  3.4 プライバシーとセキュリティ', ''),
    ('  3.5 通知と連携', ''),
    ('  3.6 紹介・招待プログラム', ''),
    ('  3.7 会計サービス連携', ''),
    ('  3.8 Apple Health 連携', ''),
    ('', ''),
    ('第 4 部 — サポート / 契約', '74'),
    ('  4.1 解約 / 返金', ''),
    ('  4.2 データのエクスポート', ''),
    ('  4.3 サポート窓口', ''),
    ('  4.4 利用規約 / プライバシーポリシー', ''),
    ('', ''),
    ('巻末資料', '78'),
    ('  用語集', ''),
    ('  改訂履歴', ''),
    ('  運営会社情報', ''),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    if title:
        r = p.add_run(title)
        r.font.size = Pt(10.5 if not title.startswith('  ') else 10)
        if not title.startswith('  '):
            r.bold = True
            r.font.color.rgb = RGBColor(0x00, 0x33, 0xA0)
        rPr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
        rPr.append(rFonts)
        if page:
            r2 = p.add_run('  ........  ' + page)
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)
page_break()


# ═══════════════════════════════════════════════════════════════════
# はじめに
# ═══════════════════════════════════════════════════════════════════
add_h1('はじめに')

add_h2('本マニュアルの使い方')
add_para(
    '本書は、株式会社 CORE が提供する「CORE Identity OS」とその 2 つのブランド (CORE Prism / CORE Iris) '
    'における全機能の操作方法・出力例・活用シーン・プラン別の利用上限を、こと細かに記述したサービス操作マニュアルです。'
)
add_para('対象読者:')
add_bullet('CORE のサービスを契約検討中の個人事業主・経営者・クリエイター')
add_bullet('すでに契約済みで、機能を網羅的に把握したいユーザー')
add_bullet('CORE を導入する代理店・販売パートナー')
add_bullet('業務システム担当者・経理担当者 (運用前提を確認したい方)')

add_para('本書の構成:')
add_para(
    '第 1 部は CORE Prism、第 2 部は CORE Iris、第 3 部は両ブランド共通機能、第 4 部はサポートや契約事項を扱います。'
    'すべての機能ブロックは下記の統一フォーマットで記載されています。'
)
add_bullet('① 何ができるか (What) — 機能の目的と概要')
add_bullet('② 出力例 / 入力例 (Example) — 実際に AI が出すもの')
add_bullet('③ 主な活用シーン (When) — どんな場面で使うか')
add_bullet('④ プラン別の使い方 (Plans) — 月の上限・含まれるプラン')

add_h2('CORE Identity OS の全体像')
add_para(
    '「CORE Identity OS」は、ひとりの人間が持つ複数の役割 (経営者・親・投資家・クリエイター 等) を、'
    '専属の AI エージェントで分散して引き受ける「自己拡張のオペレーティングシステム」です。'
    'ChatGPT のような単機能 AI ではなく、人格 (ペルソナ) 毎に独立した文脈・記憶・出力スタイルを持つ、'
    '生活/事業の OS としての位置付けです。'
)

add_callout_box('3 層構造',
    'Layer 1 (Identity)   ─ あなたの「核」(プロファイル・健康・カレンダー)\n'
    'Layer 2 (Brand)      ─ CORE Prism (事業家) / CORE Iris (創作者)\n'
    'Layer 3 (Agents)     ─ 役割ごとの専属 AI エージェント (Prism 7 / Iris 6)',
    'EAF2FF'
)

add_h2('技術スタック')
add_bullet('AI: Anthropic Claude (Opus/Sonnet/Haiku) + Google Gemini (Flash/Pro)')
add_bullet('音声: Web Speech API (入力) / OpenAI TTS (出力、6 ボイス)')
add_bullet('インフラ: Vercel Edge Functions + Supabase (Auth + Postgres + RLS)')
add_bullet('課金: Stripe Checkout + Webhook')
add_bullet('メール: Resend (トランザクション + マジックリンク)')
add_bullet('画像生成: Gemini Image API (将来 DALL-E 統合予定)')

add_h2('動作環境')
add_bullet('Web: Google Chrome / Safari / Firefox / Edge (最新版)')
add_bullet('iOS: 16.0 以降の Safari (PWA インストール推奨)')
add_bullet('Android: Chrome 110 以降 (PWA インストール推奨)')
add_bullet('macOS: Safari / Chrome (PWA でホーム追加可能)')
add_bullet('回線: 5Mbps 以上の安定したインターネット (音声機能利用時)')
page_break()


# ═══════════════════════════════════════════════════════════════════
# 第 1 部 — CORE Prism
# ═══════════════════════════════════════════════════════════════════
add_h1('第 1 部 ─── CORE Prism')

add_h2('1.1 CORE Prism とは')
add_para(
    'CORE Prism は、複数の役割を並行運用する事業家・経営者・専門職のための AI エージェント OS です。'
    'ひとつの白い光 (あなたの判断) を、7 つの人格に分散させて並列処理することを核としています。'
    'ChatGPT との最大の違いは、ペルソナごとの「継続文脈」、健康データ統合、能動提案 (デイリーブリーフ)、'
    'そして Stripe / HubSpot / Gmail 等の外部システムとの深い統合です。'
)

add_callout_box('Prism の哲学',
    '「あなたは、ひとつじゃない。経営者・営業・財務・創造者・先生・親 ── 役割の数だけ、思考が要る。」'
    '\n\nPrism は、それぞれの役割に専属の AI エージェントを配置し、文脈を分離して並行で動かします。'
    '\nあなたは指揮者になり、エージェントが演奏する ── これが Prism の世界観です。',
    'EAF2FF'
)

add_h2('1.2 7 つのエージェント')
add_para(
    '以下、7 つのエージェントを順に解説します。各エージェントは独立した記憶領域とプロンプトを持ち、'
    'ペルソナ切替時にコンテキストスイッチが発生しないよう設計されています。'
)


# ─────────────────────────────────────────────
# 1.2.1 経営エージェント
# ─────────────────────────────────────────────
add_h3('1.2.1 経営エージェント (CEO Agent)')
add_para(
    '事業全体の戦略立案・KPI モニタリング・重要な意思決定の支援を担当する、最上位のエージェント。'
    '他の 6 エージェントから上がってくる情報を統合して、「今日のあなたが判断すべきこと」を提示します。',
    italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8
)

add_h4('● 主な機能 (12 個)')
ceo_funcs = [
    '事業戦略立案ウィザード — 3-5 年のビジョンから逆算した四半期 OKR を AI と対話で作成',
    'KPI ダッシュボード — 売上 / 粗利 / CAC / LTV / Churn を Stripe API から自動取得',
    '競合分析レポート — 指定したキーワードで Web リサーチ → 3-pager レポート生成',
    '意思決定メモ (Decision Memo) — 重要決定の前後に「論点・選択肢・理由・結果」を構造化',
    'SWOT 自動更新 — 月次で実績を反映、強み/弱み/機会/脅威を再評価',
    '取締役会資料生成 — 月次の活動ハイライト + KPI を A4 サイズの PDF に',
    '事業計画書ドラフト — 投資家向けピッチデック (10 枚) を 30 分で生成',
    'M&A スクリーニング — 公開情報から条件適合候補をスクレイピング',
    'リスク棚卸し — 法務 / 財務 / 技術 / 人的 / 市場 / 評判の 6 軸でリスク評価',
    '取締役会議事録 → タスク化 — 会議録音から決議事項を抽出してカレンダー登録',
    '"今日の判断 1 件" 抽出 — 朝のブリーフで最重要決定を 1 つだけ提示',
    'ペルソナ切替時のサマリー — 他ペルソナの状況を 3 行で要約して引き継ぎ',
]
for f in ceo_funcs:
    add_bullet(f)

add_h4('● 出力例')
add_callout_box('朝のブリーフ サンプル (2026/05/12 月曜)',
    '【今日決断すべき 1 件】\n'
    '・新規顧客 A 社との契約条件、特に "契約期間 2 年 vs 1 年"。'
    '\n  → 2 年で粗利 +¥80 万、ただし違約金条項に注意。\n\n'
    '【補足: 3 件】\n'
    '① B 社の検収判定 (期限: 今日 17:00)\n'
    '② スタッフ C の 1on1 (体調変化フラグあり、慎重に)\n'
    '③ 株主への月次レポート ドラフトが用意済み (確認のみ)\n\n'
    '【身体性メモ】\n'
    '昨晩の睡眠 6.2h / HRV 38ms / 推定集中力 72%。重要交渉は午後を推奨。',
    'F2F2F7'
)

add_h4('● 活用シーン')
add_bullet('毎朝 7:00 — 一日の最重要決定を 30 秒で把握')
add_bullet('週次 KPI レビュー — 自動生成レポートを Slack に配信')
add_bullet('取締役会前夜 — 議題リスト + 各論点の AI 意見を 1 枚にまとめる')
add_bullet('競合の動向チェック — 月初に競合 3 社の四半期サマリーを自動取得')
add_bullet('重要判断の前に Decision Memo — 1 年後の自分への手紙として残す')

add_h4('● プラン別の上限')
add_para(
    'Free / Starter プランでは KPI ダッシュボードと朝のブリーフのみ利用可。Standard プラン以上で 12 機能すべて利用可能。'
    'Exclusive プランは AI モデルを Opus 4.5 (Claude 最上位) に切替え、レスポンス品質が向上します。',
    size=10, color=RGBColor(0x6E, 0x6E, 0x73)
)
add_horizontal_line()


# ─────────────────────────────────────────────
# 1.2.2 営業エージェント
# ─────────────────────────────────────────────
add_h3('1.2.2 営業エージェント (Sales Agent)')
add_para(
    'リード探索から商談スクリプト、提案書、反論対応、受注後のオンボーディング案内まで、'
    'BtoB / BtoC 問わず売上発生プロセス全体を AI が並走します。',
    italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8
)

add_h4('● 主な機能 (14 個)')
sales_funcs = [
    'リード探索 — 業界 + 地域 + 規模で公開情報から候補企業をリストアップ',
    'リード スコアリング — 過去成約データから AI が温度感 (Hot/Warm/Cold) を推定',
    '初回コンタクト メール ドラフト — 業種 + 役職に合わせた 3 案を生成',
    '商談スクリプト生成 — BANT (Budget/Authority/Need/Timing) ヒアリング項目を自動構築',
    '提案書 ドラフト (Word/PowerPoint) — 案件情報から 5-10 枚のスライドを 3 分で',
    '見積書テンプレ — 過去案件価格から AI が標準価格を提示、ボリュームディスカウントも自動計算',
    '反論対応集 — よくある懸念 (価格 / 競合 / 導入工数 / セキュリティ) ごとの回答テンプレ',
    '失注分析 — 直近 3 ヶ月の失注理由をクラスタリング、改善 3 件提案',
    'クロージング戦略 — 案件のステージ別に「次の一手」を 3 案提示',
    'フォローアップ自動化 — 商談後 N 日経過したら通知 + 内容に応じた返信ドラフト',
    'CRM 自動入力 — Gmail / Slack の会話から CRM フィールドを自動更新',
    'メーリング セグメント — 顧客リストを AI が属性別に自動セグメント',
    '商談振り返り メモ — 録音 → 文字起こし → 要点 + ネクストアクション抽出',
    'ロールプレイ — AI が顧客役として、難しい交渉のシミュレーションを実施',
]
for f in sales_funcs:
    add_bullet(f)

add_h4('● 出力例')
add_callout_box('反論対応 サンプル',
    '【顧客の反論】「同業他社の方が安いので、もう少し検討させてください」\n\n'
    '【AI ドラフト 返答 3 案】\n'
    '① 価値訴求型: 「もしお時間いただけるなら、価格ではなく "1 年後に何が変わるか" の試算をお見せできます。15 分で構いません」\n'
    '② 寄り添い型: 「価格は大切な判断軸ですよね。もしよろしければ、競合さんの見積もりを拝見できれば、純粋にどちらが御社にフィットするかをお伝えできます」\n'
    '③ 差別化型: 「弊社は ¥xx 円ですが、その中に "導入後 30 日の伴走支援" が含まれています。他社が ¥yy 円のうち、そこをご利用すると最終的に同じか弊社の方がお得になるケースが多いです」',
    'F2F2F7'
)

add_h4('● 活用シーン')
add_bullet('新規開拓週 — 月曜の朝に AI が今週リードすべき 30 社を提示')
add_bullet('商談前夜 — 顧客名 + 業界を入れるだけで、台本 + 質問項目が完成')
add_bullet('提案書作成 — 案件メモから 30 分で提案書がドラフト完了')
add_bullet('失注時 — 失注理由を入れると、改善提案が即時に出る')
add_bullet('受注後 — オンボーディングメール 5 通を AI が日数別に自動配信')

add_h4('● プラン別の上限')
add_para(
    'Starter ではリード探索 + メールドラフトのみ。Standard 以上で「商談 AI」が解放され、'
    '提案書生成・反論対応・ロールプレイなど 14 機能すべて利用可。Pro / Exclusive は CRM (HubSpot/Salesforce) との双方向同期も含まれます。',
    size=10, color=RGBColor(0x6E, 0x6E, 0x73)
)
add_horizontal_line()


# ─────────────────────────────────────────────
# 1.2.3 財務エージェント
# ─────────────────────────────────────────────
add_h3('1.2.3 財務エージェント (CFO Agent)')
add_para(
    '日次の経費入力・月次の P&L 自動生成・年次の決算ドラフトまで、'
    '中小規模事業者の財務担当者の作業を AI が代行します。会計ソフト (freee / マネーフォワード / 弥生) との双方向 OAuth 連携も可能です。',
    italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8
)

add_h4('● 主な機能 (15 個)')
cfo_funcs = [
    'レシート OCR — 写真 1 枚で日付 / 店名 / 金額 / 勘定科目を自動抽出',
    '経費精算ワークフロー — 撮影 → 仕訳 → 承認 → 振込まで一気通貫',
    '月次 P&L 自動生成 — 売上 + 経費 + 粗利 + 営業利益 を A4 1 枚に',
    'キャッシュ予測 — 直近 3 ヶ月のトレンドから 90 日先まで予測',
    '勘定科目マッピング — AI が過去仕訳から最適な科目を提案',
    '請求書 → 入金 自動消込 — Stripe / 銀行 API から消込',
    '請求書 PDF 生成 — 案件情報から国税庁適格請求書フォーマットで生成',
    '消費税自動計算 — 軽減 8% / 標準 10% を AI が判定',
    'インボイス制度対応 — 取引先の登録番号を自動チェック',
    '予算 vs 実績 ダッシュボード — 月初設定の予算と実績を比較',
    '部門別 P&L — 複数事業ある場合の部門別収益管理',
    '減価償却スケジュール — 固定資産の耐用年数を AI が判定',
    '法人税試算 — 期末予測 + 法人税額の概算',
    '監査対応資料 — 仕訳帳 / 元帳 / 残高試算表を一括出力',
    '会計士への引継ぎ書 — 月末締めデータ一式を 1 クリックで生成',
]
for f in cfo_funcs:
    add_bullet(f)

add_h4('● 出力例')
add_callout_box('月次 P&L サンプル (2026/04)',
    '【売上】 ¥3,840,000 (前月比 +12%、前年同月比 +28%)\n'
    '  ・主力商品 A: ¥2,200,000\n'
    '  ・サービス B: ¥1,640,000\n\n'
    '【経費】 ¥1,920,000\n'
    '  ・人件費: ¥1,200,000 (62.5%)\n'
    '  ・広告費: ¥320,000 (16.7%)\n'
    '  ・通信・システム: ¥180,000\n'
    '  ・その他: ¥220,000\n\n'
    '【粗利】 ¥1,920,000 (粗利率 50%)\n'
    '【営業利益】 ¥1,400,000 (営業利益率 36.5%)\n\n'
    '【AI コメント】\n'
    '広告費の ROI が前月比 -8%。月末にクリエイティブ刷新を推奨。',
    'F2F2F7'
)

add_h4('● 活用シーン')
add_bullet('毎日の経費 — レシートを撮影してアップロード、AI が即仕訳')
add_bullet('月初 5 営業日以内 — 前月 P&L が自動で配信される')
add_bullet('資金繰り会議 — 90 日先のキャッシュ予測を可視化')
add_bullet('決算前 — 仕訳の漏れや科目誤りを AI が指摘')
add_bullet('税理士への引継ぎ — クリック 1 つで月次データを引継書として出力')

add_h4('● プラン別の上限')
add_para(
    'Starter プランは月 30 件まで仕訳。Standard 以上で無制限。'
    'Pro / Exclusive は会計ソフト連携 (freee / MF / 弥生) と部門別 P&L が解放されます。'
    'なお、Phase 2 (2026 年 6 月以降) で銀行 API 直接連携 (Plaid 経由) を追加予定です。',
    size=10, color=RGBColor(0x6E, 0x6E, 0x73)
)
add_horizontal_line()

# 残りの 4 エージェントは同様のフォーマットで簡潔に
# ─── 1.2.4 創造 ───
add_h3('1.2.4 創造エージェント (Creative Agent)')
add_para('画像生成・キャプション・ブランド設計・スライド自動化など、視覚と言語の創造を担当。',
         italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8)
add_h4('● 主な機能 (13 個)')
for f in [
    '画像生成 — Gemini Image でブログ用 OG 画像、サムネ、バナーを生成',
    'キャプション作成 — Instagram / X / LinkedIn 各プラットフォーム別に最適化',
    'ブランドガイドライン作成 — 色・フォント・トーンの統一基準を自動文書化',
    'スライド自動化 — テキスト指示から 10 枚のプレゼンを 3 分で',
    'ロゴデザイン案 — 業種 + キーワードから 6 案を生成',
    'ホームページ ワイヤフレーム — 1 ページの構造を AI が設計',
    'ブログ記事 (3,000 字) — トピック指定 → SEO 最適化された記事',
    'メルマガ ドラフト — 業種 + 季節 + 顧客層から週次配信用に',
    'プレスリリース — 製品発表用の標準フォーマット',
    '会社案内 PDF — A4 8 ページの会社紹介資料',
    'CI/VI 案 — 名刺 / レターヘッド / 封筒のテンプレート',
    '動画台本 (YouTube 5 分用) — 構成 + ナレーション + テロップ案',
    'ホームページコピー — Hero / Feature / Pricing / Footer のすべてを',
]:
    add_bullet(f)
add_h4('● 出力例')
add_callout_box('キャプション (Instagram、新商品紹介、200 字)',
    '新作リップが届きました🍂 季節に寄り添うミルキーモーヴ、唇にのせた瞬間、肌のトーンがぱっと明るくなる不思議。'
    'マスクをつけても色落ちしにくく、夕方のお直し回数が半分に。今週の "整える時間" は、口元から。'
    '\n#PR #NewLip #秋メイク #モーヴリップ #大人女子コスメ',
    'F2F2F7')
add_h4('● プラン別の上限')
add_para('Starter は画像生成 月 10 枚 + スライド 3 件。Standard 以上で無制限。Pro/Exclusive はブランドガイドライン + 会社案内 PDF が含まれます。',
         size=10, color=RGBColor(0x6E, 0x6E, 0x73))
add_horizontal_line()

# ─── 1.2.5 学び ───
add_h3('1.2.5 学びエージェント (Knowledge Agent)')
add_para('読書・YouTube・ポッドキャストを要約してナレッジ化、後から横断検索 (RAG)。',
         italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8)
add_h4('● 主な機能 (11 個)')
for f in [
    'YouTube 要約 — URL を入れるだけで全動画を 5 分で要約 + 名言抽出',
    '読書ノート — 書籍タイトル + 章を入れると、重要観点を構造化',
    'ポッドキャスト書き起こし → 要約',
    'PDF / Word 取込 — 100 ページの資料も 30 秒で要約',
    '知識グラフ — 取り込んだナレッジを概念マップで可視化',
    '横断検索 (RAG) — 自然言語の質問で過去ナレッジから回答',
    'ナレッジタグ自動付与 — AI が分類タグを提案',
    '関連ナレッジ提案 — 今読んでいる内容に関連する過去ノートを提示',
    '読書計画 — 月のテーマに沿って読書リスト 5 冊を AI が提案',
    'ナレッジ 月次サマリー — 月末に今月学んだことを 1 枚にまとめる',
    'チーム共有 — Pro/Exclusive ではナレッジをチーム間で共有',
]:
    add_bullet(f)
add_h4('● プラン別の上限')
add_para('Starter は 100 件まで保存。Standard 以上で無制限。Pro/Exclusive ではチーム共有 + RAG 検索の優先度が向上。',
         size=10, color=RGBColor(0x6E, 0x6E, 0x73))
add_horizontal_line()

# ─── 1.2.6 人材 ───
add_h3('1.2.6 人材エージェント (People Agent)')
add_para('1on1 履歴・センチメント分析・採用面接準備など、ヒトに関するすべてを記憶。',
         italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8)
add_h4('● 主な機能 (12 個)')
for f in [
    '1on1 アジェンダ作成 — 過去履歴を踏まえた質問項目',
    '1on1 議事録 + ネクストアクション抽出',
    'センチメント分析 — メンバーの発言から感情推移を可視化',
    'リスク アラート — 退職リスクや疲労蓄積を早期検知',
    '採用面接ガイド — 職種別に質問項目 + 評価基準を自動生成',
    '採用候補者スコアリング — 履歴書 + 職務経歴書から AI が初期評価',
    'オファーレター ドラフト — 給与レンジ + 福利厚生を入れて生成',
    'オンボーディング計画 — 入社後 30/60/90 日の目標 + 研修プラン',
    'パフォーマンスレビュー — 半期評価のドラフト',
    '退職面談メモ — 退職理由を構造化、組織改善案を提示',
    '組織図エディタ — チーム編成変更時の影響分析',
    'メンバー紹介ポートレート — 経歴 + 強み + 写真でメンバー紹介ページ生成',
]:
    add_bullet(f)
add_h4('● プラン別の上限')
add_para('Starter は 1on1 のみ (月 5 回)。Standard 以上で全機能。Pro/Exclusive ではセンチメント分析 + リスクアラートが解放。',
         size=10, color=RGBColor(0x6E, 0x6E, 0x73))
add_horizontal_line()

# ─── 1.2.7 生活 ───
add_h3('1.2.7 生活エージェント (Life Agent)')
add_para('健康データ統合・家族のスケジュール・心の整えなど、事業以外の "人間としての時間" を担当。',
         italic=True, color=RGBColor(0x6E, 0x6E, 0x73), space_after=8)
add_h4('● 主な機能 (10 個)')
for f in [
    'Apple Health / Watch / Withings データ取込',
    '睡眠 / HRV / 心拍 / 歩数 の 60 日トレンド',
    '今日の集中力スコア (HRV ベース)',
    '家族カレンダー — 配偶者 / 子どもの予定と仕事の予定を統合',
    '記念日リマインダー — 結婚記念日 / 子どもの誕生日 / 親の誕生日',
    '健康診断結果の長期記録 + AI 解説',
    'ストレスログ — 日記形式で記録、月次傾向分析',
    '心の整え (Mindfulness) — 朝晩 3 分のガイド付き呼吸',
    '医療機関の予約管理 — 病院名 / 担当医 / 次回予約日を一元化',
    '緊急連絡先カード — 病気 / 服薬 / アレルギー情報を 1 画面に',
]:
    add_bullet(f)
add_h4('● プラン別の上限')
add_para('全プラン共通で利用可能 (健康はオプトイン、データは端末内のみ保管)。Standard 以上で家族カレンダー連動とリマインダー詳細設定。',
         size=10, color=RGBColor(0x6E, 0x6E, 0x73))
add_horizontal_line()


# ─────────────────────────────────────────────
# 1.3 Prism 横断機能
# ─────────────────────────────────────────────
add_h2('1.3 Prism 横断機能 (20 機能)')
add_para('7 つのエージェントとは別に、Prism 全体で使える機能を 20 個ご紹介します。',
         space_after=10)

cross_funcs = [
    ('AI 戦略コーチ (デイリーブリーフ)',
     '朝 (6:00-10:00) / 昼 (12:00-14:00) / 夜 (18:00-21:00) の 3 スロットで自動配信。'
     'あなたの体調・タスク・カレンダー・直近のメッセージから最も重要な判断を 3 行に圧縮。',
     'モバイルプッシュ通知 / Slack / メールでも受信可能'),
    ('議事録 AI',
     '音声録音 (Mac / iPhone マイク) → 文字起こし (Whisper) → 要約 → タスク抽出 → カレンダー登録、すべて自動。'
     '60 分の会議で 3 分の要約 + ToDo リストが完成。',
     '長時間会議 (3 時間超) は分割処理、Pro 以上で対応'),
    ('スライドスタジオ',
     '議事録 / 案件メモ / 自由テキストから 10-30 枚のスライドを自動生成。'
     'PowerPoint (.pptx) / Keynote (.key) / PDF で書き出し可。',
     'Standard 以上で 10 枚超のロング資料生成可'),
    ('書類スタジオ',
     '見積→発注→納品→請求 の 4 種類を一気通貫で生成。各書類が CRM 案件と自動連動、ステータス遷移も AI が提案。'
     '番号体系: EST/ORD/DEL/INV-{案件略号}-{年}-{連番}',
     'Standard 以上で利用可'),
    ('CRM 案件管理',
     '商談ステージ・履歴・関連書類すべてを 1 つに統合。Inquiry → Qualification → Proposal → Negotiation → Won/Lost。',
     'Pro 以上で HubSpot / Salesforce 双方向同期'),
    ('音声メモ AI 振分',
     'スマホで話すだけで 5 種類のフォルダ (アイデア / タスク / 学び / 課題 / 家族) に AI が自動振分。',
     '全プラン利用可、月 100 件まで Starter、無制限 Standard 以上'),
    ('YouTube → ナレッジ',
     'YouTube URL を貼るだけで動画要約 + 重要発言 + ナレッジベース登録まで完結。',
     '月 20 動画まで Starter、無制限 Standard 以上'),
    ('シャドー秘書',
     'Gmail を 30 分ポーリング、重要メールの返信下書きを事前生成。"確認 → 送信" の 2 タップ。',
     'Standard 以上、要 Gmail OAuth 連携'),
    ('ヘルスダッシュボード',
     'Apple Health Export (.zip / .xml) をドロップで取込。HRV / 心拍 / 睡眠 / 歩数 の 60 日トレンド。',
     '全プラン利用可、データは端末内のみ'),
    ('売上台帳 + 会計連携',
     'freee / マネーフォワード / 弥生 と双方向 OAuth で同期。請求書発行から会計入力まで自動。',
     '全プラン利用可、API キー設定後に連携'),
    ('ペルソナ切替',
     '1 タップで「経営者モード」「父親モード」「投資家モード」と文脈分離。AI は切替後のロール・記憶で動く。',
     'Starter は 1 人格のみ、Standard 以上で無制限'),
    ('音声秘書 (OpenAI TTS)',
     '6 種類のボイス (nova / alloy / echo / fable / onyx / shimmer) から選択、AI 応答を自然な音声で読み上げ。',
     'Standard 以上で利用可、月 100 分まで'),
    ('マスターモード',
     'API キー直叩きモード。Claude Opus 4.5 を無制限で利用。マスターキー "GAUCHE2026" 経由 (オーナー専用)。',
     '将来は Exclusive プランに統合予定'),
    ('能動提案パネル',
     'あなたが何もしなくても AI が「今あなたがすべきこと」を 3 つ提案。指示待ちから提案先取りへ。',
     '全プラン利用可'),
    ('アイデンティティダッシュボード',
     '7 つのエージェントの状態を 1 画面でリアルタイム可視化。指揮台に立つ感覚。',
     'Standard 以上の主画面'),
    ('PWA インストール',
     'iPhone / Mac / Android にホーム追加可。オフラインでも履歴・ナレッジは継続表示。',
     '全プラン共通、PWA は無料'),
    ('Webhook 統合',
     'Slack / Discord にデイリーブリーフを配信。複数 Webhook で複数チャネルへ同時配信。',
     '全プラン利用可'),
    ('紹介プログラム',
     '友人を紹介すると、お互いに 30 日トライアル延長。Stripe で自動精算。',
     '全プラン利用可'),
    ('招待プログラム',
     'チームメンバーや顧問をメール 1 通で招待。受信者は Magic Link クリックで参加。',
     'Pro 以上で利用可'),
    ('多言語対応',
     '日本語 / English / 中文 (簡体) の 3 言語対応。全画面で切替可。',
     '全プラン共通'),
]
for name, desc, plan in cross_funcs:
    add_h3('● ' + name)
    add_para(desc, space_after=4)
    add_para('プラン制限: ' + plan, size=9.5, color=RGBColor(0x6E, 0x6E, 0x73), italic=True, space_after=10)

page_break()


# ─────────────────────────────────────────────
# 1.4 Prism プラン詳細
# ─────────────────────────────────────────────
add_h2('1.4 CORE Prism プラン詳細')
add_para('プランは 4 段階。すべて 14 日間無料、クレジットカード登録不要で開始可能です。',
         space_after=10)

add_h3('Free — 14 日間 無料トライアル')
add_para('• 全機能を 14 日間試用可能 (制限なし)\n• クレジットカード登録不要、自動課金なし\n'
         '• 14 日経過後、自動的に閲覧専用モードへ移行 (データは保持)\n'
         '• 有料プランへのアップグレードはいつでも可能',
         size=10.5)

add_h3('Starter ¥4,980 / 月 (税抜) — 個人・スタートアップ向け')
add_para('• 基本 AI 機能 (Haiku 4.5 ベース)\n• 1 人格 / 1 ユーザー\n'
         '• ナレッジ保存 100 件まで\n• コミュニティサポート (Discord)\n'
         '• 議事録 AI 月 10 件\n• Webhook 連携 1 つ\n• 紹介プログラム有効',
         size=10.5)

add_h3('Standard ¥9,800 / 月 (税抜) — チーム本格活用 [人気 No.1]')
add_para('• 全 AI 機能 (Sonnet 4.5 ベース、商談 AI 含む)\n• 無制限人格 / 無制限ユーザー\n'
         '• ナレッジ無制限\n• OpenAI TTS 音声秘書 (月 100 分)\n'
         '• 議事録 AI 無制限\n• 書類スタジオ + CRM\n'
         '• Webhook 連携 5 つ\n• メール / Chat サポート (24 時間以内)',
         size=10.5)

add_h3('Exclusive ¥29,800 / 月 (税抜) — 経営者・大型法人向け')
add_para('• Standard 全機能 + Opus 4.5 (最高品質 AI)\n• 専任カスタマーサクセス\n'
         '• 優先サポート (1 営業日)\n• カスタム連携 (Salesforce / SAP 等)\n'
         '• 社内研修 (オンライン 2 回 + オンサイト 1 回 / 半期)\n• 導入伴走 (3 ヶ月)\n'
         '• 専有 API キー\n• SLA 99.9%',
         size=10.5)

add_h3('年払い割引')
add_para('• 年払いで 2 ヶ月分無料 (Starter ¥49,800 / 年、Standard ¥98,000 / 年、Exclusive ¥298,000 / 年)\n'
         '• 解約時は次回更新月までの分は返金なし',
         size=10.5)

page_break()


# ═══════════════════════════════════════════════════════════════════
# 第 2 部 — CORE Iris
# ═══════════════════════════════════════════════════════════════════
add_h1('第 2 部 ─── CORE Iris')

add_h2('2.1 CORE Iris とは')
add_para(
    'CORE Iris は、Instagram / YouTube / TikTok / Threads 等で発信するクリエイター・インフルエンサー・'
    'アーティスト・タレントのための AI エージェント OS です。'
    '案件管理・交渉・コンテンツ生成・コミュニティ・美容相談まで、1 つのアプリで完結します。'
    '従来の分析ツール (トレミル / SINIS 等) と異なり、"次の一手" を AI が能動的に提案する点が決定的な違いです。'
)

add_callout_box('Iris の哲学',
    '「あなたの光が、世界をつくる。」\n\n'
    'Iris = 虹彩、瞳に宿る光のかけら。'
    '\n影響力という光を、AI エージェントが見つけ、磨き、世界に放つ ── これが Iris の世界観です。'
    '\n影響力を持つほどファンとの距離が遠くなる、現代発信者の最も残酷なジレンマを解きます。',
    'FFF0F5'
)

add_h2('2.2 6 つのファセット')
add_para(
    'Iris は 6 つの角度から、創作者の活動全体をカバーします。各ファセットは独立しつつも、相互に連動して動きます。',
    space_after=10
)

# 6 ファセット 詳細
iris_facets = [
    ('2.2.1 案件 (Briefs)',
     'PR 案件・タイアップ・コラボ提案の受信から納品、レポート、入金確認までを一気通貫で管理。',
     [
         'スクショ → 案件登録 (DM 1 枚で 30 秒)',
         '案件ステージ: 打診 / 交渉中 / 受注 / 制作中 / 投稿済み / 入金待ち / 完了',
         '納期リマインダー (1 週間前 / 3 日前 / 当日)',
         'チェックリスト自動生成 (撮影 / 編集 / レビュー / 投稿 / レポート)',
         '案件別 ROI 計算 (時給換算)',
         'ブランド別過去案件履歴 (再依頼時に即座に参照)',
         '月次レポート生成 (案件数 / 売上 / カテゴリ別)',
         '確定申告サポート (源泉徴収有無 / 経費仕訳)',
     ],
     '月 10 件 (Lite) → 無制限 (Standard 以上)',
     'F3E5F5'),
    ('2.2.2 分析 (Analytics)',
     'Instagram アカウントの投稿パフォーマンス、フォロワー属性、ベストタイミングを AI が解析。',
     [
         '投稿時間最適化 (フォロワーが最もアクティブな時間を AI が推定)',
         '反応率の高いテーマ抽出 (過去 90 日のエンゲージメント分析)',
         'フォロワー属性クラスタリング (年齢 / 性別 / 興味)',
         '次の 30 日で伸びる仮説 3 つ',
         'ベンチマーク比較 (同フォロワー帯クリエイターとの比較)',
         '競合分析 (指定 3 アカウントとの差分)',
         'ハッシュタグ ROI 分析',
         'ストーリーズ視聴維持率分析',
     ],
     '月 10 回 (Standard) → 無制限 (Pro 以上)',
     'F3E5F5'),
    ('2.2.3 創作 (Creation)',
     'キャプション・サムネ・ストーリー台本・ハッシュタグ・OG 画像 を自動生成。',
     [
         '投稿キャプション (3 案を毎回提示)',
         'ストーリー台本 (15 秒 / 30 秒 / 60 秒)',
         'サムネ案 (画像 4 枚を生成)',
         'ハッシュタグ最適 30 個 (反応予測スコア付き)',
         '撮影シナリオ (商品 PR 用)',
         'リール構成 (3 シーン × 5 秒)',
         '長文ブログ 1,500 字 (note / アメブロ向け)',
         'YouTube 動画台本 (10 分用)',
     ],
     '月 30 回 (Lite) → 無制限 (Standard 以上)',
     'F3E5F5'),
    ('2.2.4 交渉 (Negotiation)',
     '料金交渉・媒体資料・ブランド提案文・カウンターオファーを AI がドラフト。',
     [
         '料金交渉ロープレ (AI が広告主役)',
         '媒体資料 PDF 自動生成 (フォロワー数 / リーチ / エンゲージメント)',
         '希望ギャラ算出 (フォロワー帯ベンチマーク + 過去案件)',
         '断り文 / カウンターオファー / 譲歩 の 3 トーン',
         '弁護士監修の利用許諾フォーマット',
         '契約書テンプレ (期間 / 範囲 / 二次利用)',
         '請求書 PDF 生成 (適格請求書対応)',
         '入金督促文 (柔らかい / 標準 / 強め の 3 トーン)',
     ],
     'Standard 以上で無制限',
     'F3E5F5'),
    ('2.2.5 ブランド (Brand)',
     'あなたの世界観 (色・フォント・トーン・撮影スタイル) を AI が学習、ブランディング資料を一括生成。',
     [
         'ブランドガイドライン PDF (10 ページ)',
         'ロゴ案 3 種',
         'フォント推奨 (見出し用 + 本文用)',
         'カラーパレット 5 色 (Primary / Secondary / Accent / Neutral / Warning)',
         '撮影スタイルガイド (構図 / 光 / 色味)',
         '名刺デザイン (両面)',
         'OG 画像 (Twitter / Facebook / LINE 用)',
         'メールヘッダ / 署名',
     ],
     'Pro 以上で利用可',
     'F3E5F5'),
    ('2.2.6 仲間 (Community)',
     '招待制クリエイターコミュニティ。案件シェア・コラボ募集・成功事例・月次研修。',
     [
         '案件シェアボード (受けきれない案件を仲間に振る)',
         'コラボ募集 (テーマ別、地域別)',
         '成功事例タイムライン (Pro 以上のメンバー限定)',
         '月次オンライン研修 (Studio 含む全員無料)',
         'オフ会情報 (季節イベント / 撮影旅行 / 勉強会)',
         '相談チャンネル (税務 / 法務 / メンタル)',
         'メンター制度 (希望者には先輩クリエイターが付く)',
         'ピアレビュー (投稿前に仲間に意見をもらえる)',
     ],
     'Standard 以上で参加可、Pro 以上で全イベント',
     'F3E5F5'),
]

for title, desc, funcs, plan_note, color in iris_facets:
    add_h3(title)
    add_para(desc, space_after=8)
    add_h4('● 主な機能')
    for f in funcs:
        add_bullet(f)
    add_h4('● プラン別の上限')
    add_para(plan_note, size=10, color=RGBColor(0x6E, 0x6E, 0x73), space_after=14)
    add_horizontal_line()


# ─────────────────────────────────────────────
# 2.3 Iris 横断機能
# ─────────────────────────────────────────────
add_h2('2.3 Iris 横断機能 (20 機能)')
add_para('6 ファセットとは別に、Iris 全体で動く機能を 20 個ご紹介します。',
         space_after=10)

iris_cross = [
    ('スクショから案件 AI 自動入力',
     'DM のスクショを 1 枚アップロードすると、ブランド名 / 報酬 / 納期 / 必須項目を 30 秒で抽出。'
     'Google Vision OCR + Gemini Pro で精度 92% (社内ベンチマーク)。',
     '月 10 回 Lite、無制限 Standard 以上'),
    ('声から投稿生成',
     '撮影帰りに 30 秒、一言吹き込むだけで、投稿文・ストーリー・ハッシュタグ・サムネ案を一括ドラフト。',
     'Standard 以上で利用可、月 30 回'),
    ('案件精査 (Triage)',
     '受信した依頼メール / DM を AI が読み、優先度 (A/B/C) / リスク / 返信案を 3 秒で提示。',
     'Standard 以上で無制限'),
    ('丸投げ編集 (Director)',
     '構成 → テロップ → キャプション → ハッシュタグ → サムネを一気にドラフト。',
     'Pro 以上で無制限、Standard は月 10 回'),
    ('美容相談 (Beauty Advisor)',
     'スキンケア・PMS・ヘア・コスメに関する相談。肌の写真を送ると AI が状態を分析、ケアプラン提示。'
     '医療判断は行わず、深刻な症状は皮膚科受診を推奨。',
     '月 50 回 Lite、無制限 Standard 以上'),
    ('Instagram 解析',
     'アカウントの過去 90 日を分析。投稿時間 / テーマ / ハッシュタグ / 競合比較。',
     '月 10 回 Standard、無制限 Pro 以上'),
    ('メディアキット PDF',
     'フォロワー数 + 過去案件 + 属性データをまとめた A4 4 ページの媒体資料を自動生成。'
     'ブランド側に送るだけで案件単価向上。',
     'Pro 以上で利用可、月 5 種類のテンプレ'),
    ('ブランドマッチ',
     'CORE Prism 側に登録された企業リストと連動。企業 ↔ クリエイターの双方向発注プラットフォーム。',
     'Pro 以上、成約時に CORE が手数料 5-10% 取得'),
    ('マルチアカウント',
     '個人 + 副ブランド + チーム、複数の Instagram アカウントをワンタップで切替。',
     'Lite 1 アカウント、Pro 5 アカウント、Studio 無制限'),
    ('ヘルス連携',
     'Apple Watch / Withings のデータと美容相談を統合。肌・睡眠・PMS 周期と投稿リズム連動。',
     '全プラン共通、データは端末内のみ'),
    ('コミュニティ',
     '招待制クリエイターコミュニティで案件・コラボ・成功事例を共有。',
     'Standard 以上で参加可'),
    ('Instagram 直接投稿',
     '生成したキャプション + 画像をワンタップで Instagram アプリへ。'
     'Web Share Level 2 + URL Scheme + クリップボードの三段戦略で OS 横断対応。',
     '全プラン共通'),
    ('音声入力',
     'Web Speech API でほぼ全機能を声で操作可能。撮影現場でハンドフリー操作。',
     '全プラン共通'),
    ('AI マネージャー',
     '「おかえり。なんでも話して。」Iris のホーム画面は 24 時間あなたの専属マネージャー。',
     '全プラン共通の主画面'),
    ('体温学習',
     'あなたの過去発信・言い回し・ニュアンスを深く学習。本物の "あなたが書いた" と区別がつかないレベル。',
     '学習開始から最低 30 投稿 (Standard 以上で常時オン)'),
    ('DM 自動応答',
     'ファンへの DM 返信を Iris が代行。挨拶 / 感想への反応 / 商品問合せ への一次返信。'
     'プロモーション送信は禁止 (Instagram 規約遵守)。',
     'Pro 以上、Instagram OAuth 連携必須'),
    ('30 日プラン (戦略アーク)',
     '次の 30 日で「何を投稿するか」を AI が戦略立案。投稿カレンダー / テーマ / 目標 KPI まで一気通貫。',
     'Standard 月 5 回、Pro 以上無制限'),
    ('ホワイトラベル (Studio)',
     '自社ブランドで Iris の機能を提供可能。事務所・代理店向けのフル カスタマイズ。'
     'ロゴ / カラー / ドメイン / 利用規約まで自社化。',
     'Studio 専用'),
    ('API 連携 (Studio)',
     'Salesforce / Slack / 任意の外部システムと API 連携。大規模事務所の業務システム統合に。',
     'Studio 専用、月 100,000 API コール'),
    ('世界観カスタム',
     '背景・カラー・フォント・絵文字すべてを、あなたの世界観に。アプリ画面そのものがブランドの延長。',
     'Pro 以上で全カスタマイズ可'),
]
for name, desc, plan in iris_cross:
    add_h3('● ' + name)
    add_para(desc, space_after=4)
    add_para('プラン制限: ' + plan, size=9.5, color=RGBColor(0x6E, 0x6E, 0x73), italic=True, space_after=10)

page_break()


# ─────────────────────────────────────────────
# 2.4 Iris プラン詳細
# ─────────────────────────────────────────────
add_h2('2.4 CORE Iris プラン詳細')
add_para('プランは 5 段階。すべて 14 日間無料、クレジットカード登録不要で開始可能です。',
         space_after=10)

add_h3('Free — 14 日間 無料トライアル')
add_para('• 全機能を 14 日間試用可能 (制限なし)\n• クレジットカード登録不要、自動課金なし\n'
         '• 14 日経過後、自動的に閲覧専用モードへ移行 (データは保持)',
         size=10.5)

add_h3('Lite ¥1,980 / 月 (税抜) — 入門・副業クリエイター')
add_para('• AI 戦略相談 月 30 回\n• 案件管理 無制限\n• スクショ AI 入力 月 10 回\n'
         '• 投稿構成・キャプション 月 30 回\n• 美容相談 月 50 回\n• 分析履歴 90 日保持\n'
         '• コミュニティ閲覧のみ',
         size=10.5)

add_h3('Standard ¥4,980 / 月 (税抜) — 本気のクリエイター [人気 No.1]')
add_para('• AI 戦略相談 無制限\n• スクショ AI 入力 無制限\n• 交渉文 AI 無制限\n'
         '• 案件精査 AI 無制限\n• 美容相談 無制限\n• 30 日プラン / 戦略アーク 月 5 回\n'
         '• Instagram 解析 月 10 回\n• 分析履歴 365 日保持\n• コミュニティ参加可',
         size=10.5)

add_h3('Pro ¥9,800 / 月 (税抜) — チーム / マネージャー')
add_para('• Standard 全機能 + 以下を追加\n• 連携アカウント 5 個まで\n• ブランドマッチ (CORE Prism 企業リスト連動)\n'
         '• 投稿カレンダー・自動 30 日プラン\n• メディアキット PDF 生成\n• 専任カスタマーサクセス (24 時間以内)\n'
         '• カスタムテンプレート (チーム共有)\n• データ無制限保存',
         size=10.5)

add_h3('Studio ¥29,800 / 月 (税抜) — 事務所 / 代理店')
add_para('• Pro 全機能 + 以下を追加\n• 連携アカウント 無制限\n• ホワイトラベル (自社ブランドで提供可)\n'
         '• API 連携 (Salesforce / Slack / 独自システム)\n• 月次オンライン研修 (代表 + メンバー 5 名分)\n'
         '• 専有 API キー\n• SLA 99.9%',
         size=10.5)

page_break()


# ═══════════════════════════════════════════════════════════════════
# 第 3 部 — 両ブランド共通機能
# ═══════════════════════════════════════════════════════════════════
add_h1('第 3 部 ─── 両ブランド共通機能')

add_h2('3.1 PWA / オフライン対応')
add_para(
    'CORE Prism / Iris ともに Progressive Web App (PWA) として実装されており、'
    'iPhone / iPad / Mac / Android の「ホーム画面に追加」でネイティブアプリと同じ操作感で利用できます。'
    'Service Worker により、オフラインでも履歴・ナレッジは継続表示されます。'
)
add_bullet('iPhone Safari → 共有 → 「ホーム画面に追加」でインストール')
add_bullet('Mac Chrome → アドレスバー右の「アプリをインストール」アイコン')
add_bullet('Android Chrome → メニュー → 「アプリをインストール」')
add_bullet('オフライン時: 履歴・ナレッジ・案件一覧は閲覧可、新規 AI 生成は不可')

add_h2('3.2 マルチテナント / 認証')
add_para(
    'Supabase Auth による Magic Link 方式を採用 (パスワード不要)。'
    'メールアドレスを入力すると認証リンクが届き、クリックするだけでログイン完了。'
    '50,000 MAU まで完全無料インフラで運用しています。'
)
add_bullet('Magic Link 認証 (パスワード漏洩リスクなし)')
add_bullet('Organization 機能 — Pro 以上でチーム招待・ロール (Owner / Admin / Member)')
add_bullet('Row Level Security による完全データ分離 (Postgres カーネルレベル)')
add_bullet('SSO / SAML 対応 (Exclusive / Studio プランで近日提供)')

add_h2('3.3 多言語対応')
add_para('日本語 / English / 中文 (簡体) の 3 言語に対応。アプリ右上の言語切替アイコンでいつでも切替可能。')
add_bullet('日本語: メインターゲット、最も翻訳精度が高い')
add_bullet('English: グローバル展開向け、US/EU 課金通貨対応')
add_bullet('中文 (簡体): 中華圏展開向け、台湾繁体は近日対応')

add_h2('3.4 プライバシーとセキュリティ')
add_para(
    'CORE は東京リージョン (asia-northeast1) のデータセンターで運用、日本国内法 (個人情報保護法) に準拠します。'
    'GDPR / CCPA にも対応しており、海外ユーザーのデータ削除要求にも即時対応します。'
)
add_bullet('データ保管: 東京リージョン (Supabase Tokyo)')
add_bullet('暗号化: 通信 TLS 1.3、保管 AES-256')
add_bullet('30 日 auto-delete: 古いブリーフ・イベントログは 30 日で自動削除')
add_bullet('ユーザー削除: マイページから 1 クリック、全データ cascade 削除')
add_bullet('医療情報: 端末内 (localStorage) のみ、サーバーには送信しない')
add_bullet('AI への送信: 必要最小限のコンテキストのみ、識別情報はマスキング')

add_h2('3.5 通知と連携')
add_bullet('Slack / Discord Webhook 連携 (デイリーブリーフ自動配信)')
add_bullet('Email 通知 (デフォルト 朝 7:00、設定で変更可)')
add_bullet('iOS Push 通知 (PWA 経由、要 OS 設定許可)')
add_bullet('Apple Continuity 連携 (Mac の通知が iPhone にも転送)')

add_h2('3.6 紹介・招待プログラム')
add_para(
    '紹介プログラム: 友人を紹介するとお互いに 30 日トライアル延長。専用リンクから登録した友人が初回課金した時点で適用。'
    '招待プログラム: チームメンバー / 顧問をメールで招待、Magic Link クリックで参加。'
)

add_h2('3.7 会計サービス連携')
add_para(
    'freee / マネーフォワード クラウド / 弥生会計オンライン の 3 サービスと双方向 OAuth 連携。'
    '請求書発行から会計入力、月次 P&L、確定申告までを 1 つのワークフローで完結します。'
)
add_bullet('OAuth 接続後、自動同期 (15 分ごと)')
add_bullet('売上自動取込 — CORE で発行した請求書が会計側に取引として記録')
add_bullet('仕訳逆送り — 会計側の仕訳が CORE の P&L レポートに反映')
add_bullet('勘定科目マッピング — 初回設定後は AI が自動判定')
add_bullet('消費税自動計算 — 軽減 / 標準を AI が判定')

add_h2('3.8 Apple Health 連携')
add_para(
    'iPhone「ヘルスケア」アプリの「すべてのヘルスケアデータを書き出す」で生成された export.zip を、'
    'CORE のヘルスダッシュボードにドロップするだけで、過去 60 日分のデータを取込できます。'
)
add_bullet('対応データ: 歩数 / アクティブカロリー / 安静時心拍 / HRV / 睡眠時間')
add_bullet('データは端末内 (localStorage) にのみ保管、CORE サーバーには送信しない')
add_bullet('Apple Watch のデータも自動的に含まれる (Watch → iPhone 同期経由)')
add_bullet('Withings / Oura / Garmin の API 連携も 2026 Q3 対応予定')

page_break()


# ═══════════════════════════════════════════════════════════════════
# 第 4 部 — サポート / 契約
# ═══════════════════════════════════════════════════════════════════
add_h1('第 4 部 ─── サポート / 契約')

add_h2('4.1 解約 / 返金')
add_para(
    'いつでも解約可能 (マイページから 1 クリック)。解約後も契約期間終了までは引き続き利用可能です。'
    '日割り返金は行いません。年払いの場合、月割り返金は次回更新月以降に適用されます。'
)
add_bullet('月払い: 当月末で解約、次月から課金停止')
add_bullet('年払い: 解約申請から次回更新月まで利用可、返金なし')
add_bullet('退会 (アカウント削除): 全データを cascade 削除、復元不可')
add_bullet('解約後 30 日間は再契約時に同じデータで復元可')

add_h2('4.2 データのエクスポート')
add_para('全データをいつでも CSV / JSON でエクスポート可能。設定 → データ管理 → エクスポート から実行。')
add_bullet('案件データ: CSV (Excel / Numbers 互換)')
add_bullet('ナレッジ: Markdown + JSON')
add_bullet('議事録: Word (.docx)')
add_bullet('チャット履歴: JSON')
add_bullet('ヘルスデータ: CSV')

add_h2('4.3 サポート窓口')
add_bullet('Email: hello@core-inc.jp')
add_bullet('Web フォーム: https://core-prism-app.vercel.app/support')
add_bullet('応答時間: 平日 24 時間以内 (Standard), 1 営業日 (Pro), 1 営業時間 (Exclusive/Studio)')
add_bullet('緊急時 (Exclusive/Studio): 専用 Slack チャンネル')

add_h2('4.4 利用規約 / プライバシーポリシー')
add_para(
    '最新版は https://core-prism-app.vercel.app/legal/terms および /legal/privacy にて公開しています。'
    '主な要点は以下の通りです。'
)
add_bullet('AI 出力の最終判断・利用責任はユーザーに帰属します')
add_bullet('医療判断 / 法的判断 / 投資助言 を行うものではありません')
add_bullet('Instagram / Twitter 等の API 利用は各プラットフォーム規約遵守')
add_bullet('生成物の著作権はユーザーに帰属、商用利用可')
add_bullet('CORE への入力データは AI 学習に二次利用しません (opt-out 不要)')

page_break()


# ═══════════════════════════════════════════════════════════════════
# 巻末資料
# ═══════════════════════════════════════════════════════════════════
add_h1('巻末資料')

add_h2('用語集')
glossary = [
    ('Core Identity OS', 'CORE が提供する上位 OS。あなたという「核」を中心に複数の役割を AI で並行運用'),
    ('ペルソナ (Persona)', 'ひとりのユーザーが持つ複数の役割 (例: 経営者 / 父 / 投資家)。AI は各ペルソナで独立した記憶を持つ'),
    ('エージェント (Agent)', 'ペルソナ内で動く専属 AI。Prism は 7 種、Iris は 6 種'),
    ('ファセット (Facet)', 'Iris における役割の分類単位。光の角度の意'),
    ('Magic Link', 'パスワード不要の認証方式。メールに届くリンクをクリックでログイン'),
    ('RLS', 'Row Level Security。Postgres カーネルレベルでテナント間データを分離'),
    ('PWA', 'Progressive Web App。ホーム画面に追加してネイティブアプリと同じ操作感'),
    ('Webhook', 'イベント発生時に外部 URL へ POST 通知する仕組み'),
    ('Decision Memo', 'CEO Agent が生成する重要決定の構造化ドキュメント'),
    ('体温学習', 'Iris がユーザーの過去発信を学習し、本人と区別がつかない発信品質を実現する機能'),
    ('30 日プラン (戦略アーク)', 'Iris の AI が次の 30 日のコンテンツ戦略を自動立案する機能'),
    ('Triage', '受信した依頼を AI が重要度別に振分ける機能'),
    ('Director', '動画 / 投稿の構成・テロップ・キャプションを一気にドラフトする機能'),
    ('ブランドマッチ', 'CORE Prism の企業リストと CORE Iris のクリエイターを双方向マッチング'),
    ('ホワイトラベル', 'Studio プランで提供する、自社ブランドで Iris の機能を提供する仕組み'),
]
for term, desc in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(term + ' ─── ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x33, 0xA0)
    r.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    for rr in [r, r2]:
        rPr = rr._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')
        rPr.append(rFonts)

add_h2('改訂履歴')
add_para('Version 1.0 — 2026/05/11 — 初版 (ベータ公開前夜)', size=10.5)
add_para('Version 1.1 — 予定 — ベータ公開後の機能追加を反映', size=10, color=RGBColor(0x6E, 0x6E, 0x73))
add_para('Version 2.0 — 予定 — マルチテナント機能の正式リリース (2026 Q3)', size=10, color=RGBColor(0x6E, 0x6E, 0x73))

add_h2('運営会社情報')
add_para('株式会社 CORE  (CORE Inc.)', bold=True, size=14)
doc.add_paragraph()
add_para('本社所在地: (法人登記準備中、2026 年 6 月設立予定)')
add_para('代表者: 井出 直毅 / Naoki Ide (Founder & CEO)')
add_para('事業内容: AI エージェント OS の開発・運営、SaaS サブスクリプション提供')
add_para('Website: https://core-prism-app.vercel.app')
add_para('Email: hello@core-inc.jp')
add_para('発行日: 2026 年 5 月 11 日')
add_para('Version: 1.0')

doc.add_paragraph()
add_horizontal_line()
add_para('本マニュアルの内容は予告なく変更されることがあります。最新版は CORE のサポートページで確認ください。',
         size=9, color=RGBColor(0x6E, 0x6E, 0x73), align='center', italic=True)
add_para('Copyright © 2026 株式会社 CORE. All rights reserved.',
         size=9, color=RGBColor(0x6E, 0x6E, 0x73), align='center')


# === 保存 ===
out_path = os.path.expanduser('~/Desktop/CORE_Identity_OS_機能説明書_2026-05-11.docx')
doc.save(out_path)
print(f'✓ Saved: {out_path}')

# 章・段落数のカウント
total_paras = sum(1 for _ in doc.paragraphs)
print(f'  Total paragraphs: {total_paras}')
